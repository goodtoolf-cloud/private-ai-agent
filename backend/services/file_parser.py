"""
FILE: backend/services/file_parser.py
Parse uploaded files (PDF, DOCX, Excel, CSV, images, audio).
"""

import os
import io
from pathlib import Path
import PyPDF2
import polars as pd
from docx import Document as DocxDocument
import base64
from PIL import Image


def parse_file(filepath: str, file_type: str) -> dict:
    """
    Parse any supported file and return extracted content.
    Returns: { "text": str, "tables": list, "images": list, "metadata": dict }
    """
    ext = Path(filepath).suffix.lower().lstrip(".")
    if file_type:
        ext = file_type.lower()

    if ext == "pdf":
        return _parse_pdf(filepath)
    elif ext in ("docx", "doc"):
        return _parse_docx(filepath)
    elif ext in ("xlsx", "xls"):
        return _parse_excel(filepath)
    elif ext == "csv":
        return _parse_csv(filepath)
    elif ext in ("png", "jpg", "jpeg", "gif", "bmp", "webp"):
        return _parse_image(filepath)
    elif ext in ("txt", "md", "json", "xml", "html", "py", "js", "ts", "yaml", "yml"):
        return _parse_text(filepath)
    else:
        return {"text": f"Unsupported file type: {ext}", "tables": [], "images": [], "metadata": {}}


def _parse_pdf(filepath: str) -> dict:
    text_parts = []
    metadata = {}
    try:
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            metadata = {
                "pages": len(reader.pages),
                "title": reader.metadata.get("/Title", "") if reader.metadata else "",
                "author": reader.metadata.get("/Author", "") if reader.metadata else "",
            }
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
    except Exception as e:
        return {"text": f"PDF parse error: {e}", "tables": [], "images": [], "metadata": {}}

    return {
        "text": "\n\n".join(text_parts),
        "tables": [],
        "images": [],
        "metadata": metadata,
    }


def _parse_docx(filepath: str) -> dict:
    try:
        doc = DocxDocument(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append(rows)
        return {
            "text": "\n".join(paragraphs),
            "tables": tables,
            "images": [],
            "metadata": {"paragraphs": len(paragraphs), "tables": len(tables)},
        }
    except Exception as e:
        return {"text": f"DOCX parse error: {e}", "tables": [], "images": [], "metadata": {}}


def _parse_excel(filepath: str) -> dict:
    try:
        xf = pd.ExcelFile(filepath)
        all_text = []
        all_tables = []
        for sheet in xf.sheet_names:
            df = pd.read_excel(filepath, sheet_name=sheet)
            all_text.append(f"Sheet: {sheet}\n{df.to_string(index=False)}")
            all_tables.append(df.values.tolist())
        return {
            "text": "\n\n".join(all_text),
            "tables": all_tables,
            "images": [],
            "metadata": {"sheets": xf.sheet_names},
        }
    except Exception as e:
        return {"text": f"Excel parse error: {e}", "tables": [], "images": [], "metadata": {}}


def _parse_csv(filepath: str) -> dict:
    try:
        df = pd.read_csv(filepath)
        return {
            "text": df.to_string(index=False),
            "tables": [df.values.tolist()],
            "images": [],
            "metadata": {"rows": len(df), "columns": list(df.columns)},
        }
    except Exception as e:
        return {"text": f"CSV parse error: {e}", "tables": [], "images": [], "metadata": {}}


def _parse_image(filepath: str) -> dict:
    try:
        img = Image.open(filepath)
        metadata = {
            "width": img.width,
            "height": img.height,
            "mode": img.mode,
            "format": img.format,
        }
        # Encode as base64 for multimodal models
        with open(filepath, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("utf-8")
        return {
            "text": f"[Image: {img.width}x{img.height} {img.format}]",
            "tables": [],
            "images": [b64],
            "metadata": metadata,
        }
    except Exception as e:
        return {"text": f"Image parse error: {e}", "tables": [], "images": [], "metadata": {}}


def _parse_text(filepath: str) -> dict:
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        return {
            "text": content,
            "tables": [],
            "images": [],
            "metadata": {"chars": len(content), "lines": content.count("\n")},
        }
    except Exception as e:
        return {"text": f"Text parse error: {e}", "tables": [], "images": [], "metadata": {}}
