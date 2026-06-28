"""
FILE: backend/services/document_creator.py
Create Word, PDF, and Excel files programmatically.
"""

import os
import uuid
from io import BytesIO
from config import get_settings

settings = get_settings()


def _out_path(filename: str) -> str:
    path = os.path.join(settings.generated_dir, "documents", filename)
    os.makedirs(os.path.dirname(path), exist_ok=True)
    return path


async def create_docx(title: str, content: str, sections: list[dict] = []) -> dict:
    """
    Create a .docx Word document.
    sections: [{"heading": str, "body": str}, ...]
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # If sections provided, use them; otherwise use raw content
    if sections:
        for sec in sections:
            if sec.get("heading"):
                doc.add_heading(sec["heading"], level=1)
            if sec.get("body"):
                doc.add_paragraph(sec["body"])
    else:
        for line in content.split("\n"):
            if line.startswith("## "):
                doc.add_heading(line[3:], level=1)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=2)
            elif line.strip():
                doc.add_paragraph(line)

    filename = f"{uuid.uuid4()}.docx"
    out = _out_path(filename)
    doc.save(out)

    return {"filename": filename, "url": f"/generated/documents/{filename}", "type": "docx"}


async def create_pdf(title: str, content: str) -> dict:
    """Create a .pdf document."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import cm
    from reportlab.lib import colors

    filename = f"{uuid.uuid4()}.pdf"
    out = _out_path(filename)

    doc = SimpleDocTemplate(out, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontSize=20,
        spaceAfter=20,
        textColor=colors.HexColor("#1e3a5f"),
    )
    heading_style = ParagraphStyle(
        "CustomHeading",
        parent=styles["Heading1"],
        fontSize=14,
        spaceAfter=10,
        textColor=colors.HexColor("#2d6a4f"),
    )
    body_style = styles["BodyText"]

    story = [Paragraph(title, title_style), Spacer(1, 0.5 * cm)]

    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 0.3 * cm))
        elif stripped.startswith("## "):
            story.append(Paragraph(stripped[3:], heading_style))
        elif stripped.startswith("# "):
            story.append(Paragraph(stripped[2:], heading_style))
        else:
            story.append(Paragraph(stripped, body_style))

    doc.build(story)
    return {"filename": filename, "url": f"/generated/documents/{filename}", "type": "pdf"}


async def create_xlsx(title: str, data: list[list], headers: list[str] = []) -> dict:
    """Create an .xlsx Excel file from data rows."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = title[:31]

    header_fill = PatternFill("solid", fgColor="1e3a5f")
    header_font = Font(color="FFFFFF", bold=True)

    if headers:
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center")

    start_row = 2 if headers else 1
    for row_idx, row in enumerate(data, start_row):
        for col_idx, val in enumerate(row, 1):
            ws.cell(row=row_idx, column=col_idx, value=val)

    # Auto-width columns
    for col in ws.columns:
        max_len = max((len(str(cell.value or "")) for cell in col), default=0)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)

    filename = f"{uuid.uuid4()}.xlsx"
    out = _out_path(filename)
    wb.save(out)

    return {"filename": filename, "url": f"/generated/documents/{filename}", "type": "xlsx"}
